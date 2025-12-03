"""
Unit tests for DOCXReportGenerator.

Tests cover:
- Basic DOCX generation
- Configuration validation
- Data format support (DataFrame, dict, list)
- Metadata inclusion
- Table alignment
- Custom fonts
- Page breaks
- Table of contents placeholder
- Supported features
- Error handling
"""

from pathlib import Path

import pandas as pd
import pytest
from docx import Document  # For DOCX validation

from extract_transform_platform.reports import (
    DOCXReportConfig,
    DOCXReportGenerator,
)


class TestDOCXReportGenerator:
    """Test suite for DOCXReportGenerator."""

    @pytest.fixture
    def sample_data(self) -> pd.DataFrame:
        """Sample employee data for testing."""
        return pd.DataFrame(
            {
                "Employee": ["Alice Johnson", "Bob Smith", "Charlie Brown"],
                "Department": ["Engineering", "Sales", "Marketing"],
                "Salary": [95000, 85000, 75000],
                "Start Date": ["2020-01-15", "2019-06-01", "2021-03-10"],
            }
        )

    @pytest.fixture
    def output_path(self, tmp_path: Path) -> Path:
        """Temporary output path for test reports."""
        return tmp_path / "test_report.docx"

    def test_initialization(self) -> None:
        """Test generator initialization."""
        generator = DOCXReportGenerator()
        assert generator is not None
        assert len(generator.get_supported_features()) > 0

    def test_generate_basic_docx(
        self, sample_data: pd.DataFrame, output_path: Path
    ) -> None:
        """Test basic DOCX generation with default config."""
        generator = DOCXReportGenerator()
        config = DOCXReportConfig(title="Employee Report")

        result_path = generator.generate(sample_data, output_path, config)

        # Verify file exists
        assert result_path.exists()
        assert result_path == output_path.resolve()

        # Verify file is valid DOCX
        doc = Document(result_path)
        assert len(doc.tables) == 1  # One table

        # Verify table dimensions (header + 3 data rows)
        table = doc.tables[0]
        assert len(table.rows) == 4  # Header + 3 data rows
        assert len(table.columns) == 4  # 4 columns

        # Verify header row content
        header_cells = table.rows[0].cells
        assert header_cells[0].text == "Employee"
        assert header_cells[1].text == "Department"
        assert header_cells[2].text == "Salary"
        assert header_cells[3].text == "Start Date"

        # Verify first data row
        first_row_cells = table.rows[1].cells
        assert first_row_cells[0].text == "Alice Johnson"
        assert first_row_cells[1].text == "Engineering"
        assert first_row_cells[2].text == "95000"

    def test_docx_with_metadata(
        self, sample_data: pd.DataFrame, output_path: Path
    ) -> None:
        """Test DOCX with timestamp metadata."""
        config = DOCXReportConfig(
            title="Employee Report with Metadata", include_timestamp=True
        )
        generator = DOCXReportGenerator()
        result_path = generator.generate(sample_data, output_path, config)

        assert result_path.exists()
        doc = Document(result_path)

        # Check metadata paragraph exists (contains "Generated" keyword)
        paragraphs_text = [p.text for p in doc.paragraphs]
        assert any("Generated" in text for text in paragraphs_text)

    def test_docx_without_metadata(
        self, sample_data: pd.DataFrame, output_path: Path
    ) -> None:
        """Test DOCX without metadata."""
        config = DOCXReportConfig(title="Employee Report", include_timestamp=False)
        generator = DOCXReportGenerator()
        result_path = generator.generate(sample_data, output_path, config)

        assert result_path.exists()
        doc = Document(result_path)

        # Check metadata paragraph doesn't exist
        paragraphs_text = [p.text for p in doc.paragraphs]
        assert not any("Generated" in text for text in paragraphs_text)

    def test_docx_table_alignment_left(
        self, sample_data: pd.DataFrame, tmp_path: Path
    ) -> None:
        """Test left table alignment."""
        config = DOCXReportConfig(title="Left Aligned Report", table_alignment="left")
        generator = DOCXReportGenerator()
        output_path = tmp_path / "test_left.docx"
        result_path = generator.generate(sample_data, output_path, config)

        assert result_path.exists()
        doc = Document(result_path)
        assert len(doc.tables) == 1

    def test_docx_table_alignment_center(
        self, sample_data: pd.DataFrame, tmp_path: Path
    ) -> None:
        """Test center table alignment."""
        config = DOCXReportConfig(
            title="Center Aligned Report", table_alignment="center"
        )
        generator = DOCXReportGenerator()
        output_path = tmp_path / "test_center.docx"
        result_path = generator.generate(sample_data, output_path, config)

        assert result_path.exists()
        doc = Document(result_path)
        assert len(doc.tables) == 1

    def test_docx_table_alignment_right(
        self, sample_data: pd.DataFrame, tmp_path: Path
    ) -> None:
        """Test right table alignment."""
        config = DOCXReportConfig(title="Right Aligned Report", table_alignment="right")
        generator = DOCXReportGenerator()
        output_path = tmp_path / "test_right.docx"
        result_path = generator.generate(sample_data, output_path, config)

        assert result_path.exists()
        doc = Document(result_path)
        assert len(doc.tables) == 1

    def test_docx_custom_font(
        self, sample_data: pd.DataFrame, output_path: Path
    ) -> None:
        """Test custom font settings."""
        config = DOCXReportConfig(
            title="Custom Font Report", font_name="Arial", font_size=12
        )
        generator = DOCXReportGenerator()
        result_path = generator.generate(sample_data, output_path, config)

        assert result_path.exists()
        doc = Document(result_path)
        # Verify document created successfully
        assert len(doc.tables) == 1

    def test_docx_with_page_break(
        self, sample_data: pd.DataFrame, output_path: Path
    ) -> None:
        """Test page break after title."""
        config = DOCXReportConfig(
            title="Report with Page Break", page_break_after_title=True
        )
        generator = DOCXReportGenerator()
        result_path = generator.generate(sample_data, output_path, config)

        assert result_path.exists()
        doc = Document(result_path)
        # Verify document created successfully
        assert len(doc.tables) == 1

    def test_docx_with_toc(self, sample_data: pd.DataFrame, output_path: Path) -> None:
        """Test table of contents placeholder."""
        config = DOCXReportConfig(title="Report with TOC", include_toc=True)
        generator = DOCXReportGenerator()
        result_path = generator.generate(sample_data, output_path, config)

        assert result_path.exists()
        doc = Document(result_path)

        # Check TOC placeholder exists
        paragraphs_text = [p.text for p in doc.paragraphs]
        assert any("Table of Contents" in text for text in paragraphs_text)

    def test_docx_heading_levels(
        self, sample_data: pd.DataFrame, tmp_path: Path
    ) -> None:
        """Test different heading levels."""
        for level in [1, 2, 3, 5, 9]:
            config = DOCXReportConfig(
                title=f"Level {level} Heading", heading_level=level
            )
            generator = DOCXReportGenerator()
            output_path = tmp_path / f"test_heading_{level}.docx"
            result_path = generator.generate(sample_data, output_path, config)

            assert result_path.exists()
            doc = Document(result_path)
            assert len(doc.tables) == 1

    def test_docx_custom_table_style(
        self, sample_data: pd.DataFrame, output_path: Path
    ) -> None:
        """Test custom table style."""
        config = DOCXReportConfig(
            title="Custom Style Report", table_style="Medium Grid 1 Accent 1"
        )
        generator = DOCXReportGenerator()
        result_path = generator.generate(sample_data, output_path, config)

        assert result_path.exists()
        doc = Document(result_path)
        assert len(doc.tables) == 1

    def test_get_supported_features(self) -> None:
        """Test supported features list."""
        generator = DOCXReportGenerator()
        features = generator.get_supported_features()

        # Verify expected features
        assert "tables" in features
        assert "headings" in features
        assert "paragraphs" in features
        assert "styles" in features
        assert "table_of_contents" in features
        assert "page_breaks" in features
        assert "alignment" in features

        # Verify at least 7 features
        assert len(features) >= 7

    def test_invalid_output_extension_raises(
        self, sample_data: pd.DataFrame, tmp_path: Path
    ) -> None:
        """Test invalid extension raises error."""
        generator = DOCXReportGenerator()
        config = DOCXReportConfig(title="Test")
        invalid_path = tmp_path / "test.txt"

        with pytest.raises(ValueError, match="must have .docx extension"):
            generator.generate(sample_data, invalid_path, config)

    def test_wrong_config_type_raises(
        self, sample_data: pd.DataFrame, output_path: Path
    ) -> None:
        """Test wrong config type raises error."""
        from extract_transform_platform.reports import ExcelReportConfig

        generator = DOCXReportGenerator()
        wrong_config = ExcelReportConfig(title="Test")

        with pytest.raises(TypeError, match="requires DOCXReportConfig"):
            generator.generate(sample_data, output_path, wrong_config)

    def test_dict_input(self, tmp_path: Path) -> None:
        """Test single dict input."""
        generator = DOCXReportGenerator()
        config = DOCXReportConfig(title="Dict Input Report")
        data = {"Name": "Alice", "Age": 30, "City": "NYC"}
        output_path = tmp_path / "dict_test.docx"

        result_path = generator.generate(data, output_path, config)
        assert result_path.exists()

        doc = Document(result_path)
        table = doc.tables[0]
        assert len(table.rows) == 2  # Header + 1 data row
        assert len(table.columns) == 3  # 3 columns

    def test_list_of_dicts_input(self, tmp_path: Path) -> None:
        """Test list of dicts input."""
        generator = DOCXReportGenerator()
        config = DOCXReportConfig(title="List Input Report")
        data = [{"Name": "Alice", "Age": 30}, {"Name": "Bob", "Age": 25}]
        output_path = tmp_path / "list_test.docx"

        result_path = generator.generate(data, output_path, config)
        assert result_path.exists()

        doc = Document(result_path)
        table = doc.tables[0]
        assert len(table.rows) == 3  # Header + 2 data rows
        assert len(table.columns) == 2  # 2 columns

    def test_empty_dataframe_raises(self, output_path: Path) -> None:
        """Test empty DataFrame raises error."""
        generator = DOCXReportGenerator()
        config = DOCXReportConfig(title="Empty Test")
        empty_df = pd.DataFrame()

        with pytest.raises(ValueError, match="cannot be empty"):
            generator.generate(empty_df, output_path, config)

    def test_none_data_raises(self, output_path: Path) -> None:
        """Test None data raises error."""
        generator = DOCXReportGenerator()
        config = DOCXReportConfig(title="None Test")

        with pytest.raises(ValueError, match="cannot be None"):
            generator.generate(None, output_path, config)

    def test_unsupported_data_type_raises(self, output_path: Path) -> None:
        """Test unsupported data type raises error."""
        generator = DOCXReportGenerator()
        config = DOCXReportConfig(title="Invalid Type Test")
        invalid_data = "not a valid type"

        with pytest.raises(TypeError, match="Unsupported data type"):
            generator.generate(invalid_data, output_path, config)

    def test_large_dataset(self, tmp_path: Path) -> None:
        """Test with larger dataset (performance check)."""
        # Create 1000 row dataset
        large_data = pd.DataFrame(
            {
                "ID": range(1000),
                "Name": [f"Person {i}" for i in range(1000)],
                "Value": [i * 100 for i in range(1000)],
            }
        )

        generator = DOCXReportGenerator()
        config = DOCXReportConfig(title="Large Dataset Report")
        output_path = tmp_path / "large_test.docx"

        result_path = generator.generate(large_data, output_path, config)
        assert result_path.exists()

        doc = Document(result_path)
        table = doc.tables[0]
        assert len(table.rows) == 1001  # Header + 1000 data rows
        assert len(table.columns) == 3

    def test_output_path_creates_parent_directories(
        self, sample_data: pd.DataFrame, tmp_path: Path
    ) -> None:
        """Test that parent directories are created automatically."""
        nested_path = tmp_path / "level1" / "level2" / "report.docx"
        generator = DOCXReportGenerator()
        config = DOCXReportConfig(title="Nested Path Test")

        result_path = generator.generate(sample_data, nested_path, config)
        assert result_path.exists()
        assert result_path.parent.exists()

    def test_docx_file_size_reasonable(
        self, sample_data: pd.DataFrame, output_path: Path
    ) -> None:
        """Test that generated file size is reasonable."""
        generator = DOCXReportGenerator()
        config = DOCXReportConfig(title="File Size Test")

        result_path = generator.generate(sample_data, output_path, config)
        file_size = result_path.stat().st_size

        # Small dataset should produce file less than 50KB
        assert file_size < 50_000

    def test_config_validation(self) -> None:
        """Test config field validation."""
        # Valid config
        valid_config = DOCXReportConfig(
            title="Test", heading_level=5, font_size=12, table_alignment="center"
        )
        assert valid_config.heading_level == 5
        assert valid_config.font_size == 12
        assert valid_config.table_alignment == "center"

        # Invalid heading level (too high)
        with pytest.raises(Exception):  # Pydantic ValidationError
            DOCXReportConfig(title="Test", heading_level=10)

        # Invalid heading level (too low)
        with pytest.raises(Exception):  # Pydantic ValidationError
            DOCXReportConfig(title="Test", heading_level=0)

        # Invalid table alignment
        with pytest.raises(Exception):  # Pydantic ValidationError
            DOCXReportConfig(title="Test", table_alignment="invalid")
